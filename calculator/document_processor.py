"""
Document processing module for PDF, DOCX, and image files
Extracts text, creates embeddings, and extracts features
"""

import os
import json
import uuid
from pathlib import Path
from typing import Dict, List, Optional, Tuple

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import pytesseract
except ImportError:
    pytesseract = None

from django.conf import settings


class DocumentProcessor:
    """Process different document types and extract text"""

    DOCUMENTS_DIR = Path(settings.BASE_DIR) / 'documents_storage'
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    SUPPORTED_TYPES = ['pdf', 'docx', 'image', 'txt']

    @staticmethod
    def ensure_storage_dir():
        """Ensure documents_storage directory exists"""
        DocumentProcessor.DOCUMENTS_DIR.mkdir(exist_ok=True)

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> Tuple[str, int]:
        """
        Extract text from PDF file

        Args:
            file_path: Full path to PDF file

        Returns:
            Tuple of (extracted_text, page_count)
        """
        if PyPDF2 is None:
            raise ImportError("PyPDF2 is not installed. Install with: pip install PyPDF2")

        try:
            text_content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)

                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text:
                            text_content.append(text)
                    except Exception as e:
                        print(f"Error extracting page {page_num + 1}: {e}")

            extracted_text = '\n'.join(text_content)
            return extracted_text, page_count

        except Exception as e:
            raise ValueError(f"Error processing PDF: {str(e)}")

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """
        Extract text from DOCX file

        Args:
            file_path: Full path to DOCX file

        Returns:
            Extracted text content
        """
        if DocxDocument is None:
            raise ImportError("python-docx is not installed. Install with: pip install python-docx")

        try:
            doc = DocxDocument(file_path)
            text_content = []

            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    if any(row_text):
                        text_content.append(' | '.join(row_text))

            extracted_text = '\n'.join(text_content)
            return extracted_text

        except Exception as e:
            raise ValueError(f"Error processing DOCX: {str(e)}")

    @staticmethod
    def extract_text_from_image(file_path: str) -> Tuple[str, int, int, str]:
        """
        Extract text from image using OCR (Tesseract)

        Args:
            file_path: Full path to image file

        Returns:
            Tuple of (extracted_text, width, height, image_format)
        """
        if Image is None:
            raise ImportError("Pillow is not installed. Install with: pip install Pillow")

        if pytesseract is None:
            raise ImportError("pytesseract is not installed. Install with: pip install pytesseract")

        try:
            img = Image.open(file_path)
            width, height = img.size
            image_format = img.format or 'unknown'

            # Extract text using OCR
            text = pytesseract.image_to_string(img, lang='deu+eng')

            return text, width, height, image_format

        except Exception as e:
            raise ValueError(f"Error processing image: {str(e)}")

    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """
        Read text from plain text file

        Args:
            file_path: Full path to text file

        Returns:
            File content
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            raise ValueError(f"Error reading text file: {str(e)}")

    @staticmethod
    def process_document(file_path: str, file_type: str) -> Dict:
        """
        Process a document based on its type

        Args:
            file_path: Full path to document
            file_type: Type of document (pdf, docx, image, txt)

        Returns:
            Dictionary with extracted content and metadata
        """
        result = {
            'text_content': '',
            'metadata': {},
            'success': False,
            'error': None
        }

        try:
            if file_type == 'pdf':
                text, pages = DocumentProcessor.extract_text_from_pdf(file_path)
                result['text_content'] = text
                result['metadata']['page_count'] = pages

            elif file_type == 'docx':
                text = DocumentProcessor.extract_text_from_docx(file_path)
                result['text_content'] = text

            elif file_type == 'image':
                text, width, height, img_format = DocumentProcessor.extract_text_from_image(file_path)
                result['text_content'] = text
                result['metadata']['width'] = width
                result['metadata']['height'] = height
                result['metadata']['format'] = img_format

            elif file_type == 'txt':
                text = DocumentProcessor.extract_text_from_txt(file_path)
                result['text_content'] = text

            result['success'] = True

        except Exception as e:
            result['error'] = str(e)
            result['success'] = False

        return result

    @staticmethod
    def save_uploaded_file(uploaded_file, original_filename: str) -> str:
        """
        Save uploaded file to documents_storage

        Args:
            uploaded_file: Django uploaded file object
            original_filename: Original filename

        Returns:
            Relative file path
        """
        DocumentProcessor.ensure_storage_dir()

        # Generate unique filename
        file_ext = Path(original_filename).suffix
        unique_filename = f"{uuid.uuid4()}{file_ext}"
        rel_path = f"{unique_filename}"
        full_path = DocumentProcessor.DOCUMENTS_DIR / unique_filename

        # Save file
        with open(full_path, 'wb') as destination:
            for chunk in uploaded_file.chunks():
                destination.write(chunk)

        return rel_path

    @staticmethod
    def get_file_path(rel_path: str) -> str:
        """Get full file path from relative path"""
        return str(DocumentProcessor.DOCUMENTS_DIR / rel_path)

    @staticmethod
    def get_file_size(file_path: str) -> int:
        """Get file size in bytes"""
        try:
            return os.path.getsize(file_path)
        except:
            return 0

    @staticmethod
    def create_text_preview(text: str, max_chars: int = 500) -> str:
        """Create preview from text content"""
        if not text:
            return ""
        preview = text[:max_chars]
        if len(text) > max_chars:
            preview += "..."
        return preview


class DocumentEmbedder:
    """Create embeddings for documents using Sentence Transformers"""

    def __init__(self):
        try:
            from sentence_transformers import SentenceTransformer
            self.model = SentenceTransformer('sentence-transformers/paraphrase-multilingual-mpnet-base-v2')
        except ImportError:
            raise ImportError("sentence-transformers not installed")

    def embed_text(self, text: str) -> List[float]:
        """
        Create embedding for text

        Args:
            text: Text to embed

        Returns:
            Vector embedding (list of floats)
        """
        if not text or not text.strip():
            return None

        try:
            # Truncate if too long
            text = text[:1000]
            embedding = self.model.encode(text)
            return embedding.tolist()
        except Exception as e:
            print(f"Error creating embedding: {e}")
            return None


class DocumentFeatureExtractor:
    """Extract features from documents for ML model"""

    @staticmethod
    def extract_features(document_text: str, file_type: str, metadata: Dict) -> Dict:
        """
        Extract features from document

        Args:
            document_text: Extracted text content
            file_type: Type of document
            metadata: Document metadata

        Returns:
            Dictionary of extracted features
        """
        features = {
            'text_length': len(document_text) if document_text else 0,
            'word_count': len(document_text.split()) if document_text else 0,
            'file_type': file_type,
        }

        # Type-specific features
        if file_type == 'pdf':
            features['page_count'] = metadata.get('page_count', 0)
            features['avg_words_per_page'] = (
                features['word_count'] / features['page_count']
                if features['page_count'] > 0 else 0
            )

        elif file_type == 'image':
            features['image_width'] = metadata.get('width', 0)
            features['image_height'] = metadata.get('height', 0)
            features['image_area'] = features['image_width'] * features['image_height']

        # Content analysis
        if document_text:
            text_lower = document_text.lower()
            features['has_pricing_info'] = any(
                word in text_lower for word in ['preis', 'euro', '€', 'kosten', 'gebühr', 'tarif']
            )
            features['has_project_info'] = any(
                word in text_lower for word in ['projekt', 'baustelle', 'auftrag', 'arbeiten']
            )
            features['has_material_info'] = any(
                word in text_lower for word in ['material', 'holz', 'farbe', 'lack', 'stoff']
            )

        return features


class DocumentSearcher:
    """Full-text search for documents"""

    @staticmethod
    def prepare_search_text(document_text: str) -> str:
        """Prepare text for full-text search (lowercase, clean)"""
        if not document_text:
            return ""

        # Convert to lowercase
        text = document_text.lower()

        # Remove extra whitespace
        text = ' '.join(text.split())

        return text

    @staticmethod
    def search_documents(query: str, documents: List) -> List[Tuple[str, float]]:
        """
        Search documents using simple text matching

        Args:
            query: Search query
            documents: List of document objects

        Returns:
            List of (document_id, relevance_score) tuples
        """
        query_lower = query.lower()
        results = []

        for doc in documents:
            if not doc.searchable_text:
                continue

            # Count occurrences of query terms
            relevance = 0
            for term in query_lower.split():
                # Simple relevance: count matches
                relevance += doc.searchable_text.count(term)

            if relevance > 0:
                results.append((doc.id, relevance))

        # Sort by relevance (descending)
        results.sort(key=lambda x: x[1], reverse=True)

        return results


# Test if modules are installed
def check_dependencies():
    """Check if all optional dependencies are installed"""
    status = {
        'PyPDF2': PyPDF2 is not None,
        'python-docx': DocxDocument is not None,
        'Pillow': Image is not None,
        'pytesseract': pytesseract is not None,
        'sentence-transformers': True  # Already installed in Phase 1
    }
    return status
